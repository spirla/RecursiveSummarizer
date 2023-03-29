import pandas as pd
import openai
import docx

# Lê o arquivo do Excel
df = pd.read_excel('tabela.xlsx')

# Inicializa o OpenAI
openai.api_key = open_file('openaiapikey.txt')

# Cria um documento do Word
document = docx.Document()

# Loop pelas linhas da tabela
for index, row in df.iterrows():
    pergunta = row['Perguntas']
    resposta = row['Respostas']
    
    # Faz uma requisição para a API do OpenAI
    completions = openai.Completion.create(
        engine="text-davinci-003",
        prompt=resposta,
        max_tokens=2000,
        n=1,
        stop=None,
        temperature=0.5,
    )
    
    # Adiciona a pergunta e a resposta melhorada ao documento do Word
    document.add_paragraph(pergunta)
    document.add_paragraph(completions.choices[0].text)
    document.add_paragraph("***")

# Salva o documento do Word
document.save('respostas_melhoradas.docx')
